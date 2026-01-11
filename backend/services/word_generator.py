"""Word document report generator service."""
from io import BytesIO
from typing import Dict, Any

from docx import Document


def generate_word_report(report_data: Dict[str, Any]) -> BytesIO:
    """Generate a Word document report for a pupil."""
    doc = Document()

    # Title
    title = doc.add_heading(f"Development Report: {report_data['pupil_name']}", 0)

    # Period
    period = f"Period: {report_data['start_date']} - {report_data['end_date']}"
    doc.add_paragraph(period)

    # Entries by category
    for category, entries in report_data.get('entries_by_category', {}).items():
        doc.add_heading(category, level=1)
        for entry in entries:
            entry_text = f"[{entry['date']}] {entry['text']}"
            if entry.get('grade'):
                entry_text += f" (Grade: {entry['grade']})"
            doc.add_paragraph(entry_text, style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
